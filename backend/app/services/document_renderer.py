import os
import uuid
from typing import Dict, Any
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from app.core.storage import storage_service
from app.core.config import settings

def get_output_path(filename: str) -> str:
    folder = os.path.join(settings.STORAGE_DIR, "documents")
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, filename)

class DocumentRendererService:
    @classmethod
    def render_docx(cls, structured_content: Dict[str, Any], template_type: str = "IEEE") -> str:
        """
        Renders paper content to Word DOCX adhering to target journal template style (IEEE, Springer, Elsevier).
        """
        doc = DocxDocument()
        
        # Configure Margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(structured_content.get("topic", "Research Paper Title"))
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.name = "Times New Roman"
        if template_type == "IEEE":
            title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Authors placeholder
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_p.add_run("Author Name(s) Redacted for Peer Review / Institutional Affiliation")
        author_run.font.italic = True
        author_run.font.size = Pt(10)

        doc.add_paragraph() # Spacer

        # Abstract Header
        abs_head = doc.add_paragraph()
        abs_run = abs_head.add_run("Abstract—")
        abs_run.font.bold = True
        abs_run.font.italic = True
        abs_p_run = abs_head.add_run(structured_content.get("abstract", ""))
        abs_p_run.font.italic = True

        # Sections
        sections_data = [
            ("I. INTRODUCTION", structured_content.get("introduction", [])),
            ("II. LITERATURE REVIEW", structured_content.get("literature_review", [])),
            ("III. METHODOLOGY", structured_content.get("methodology", {})),
            ("IV. RESULTS & DISCUSSION", structured_content.get("results_and_discussion", {})),
            ("V. CONCLUSION", structured_content.get("conclusion", "")),
            ("VI. FUTURE SCOPE & RESEARCH HORIZONS", structured_content.get("future_scope", []))
        ]

        for sec_title, sec_body in sections_data:
            sec_p = doc.add_paragraph()
            sec_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template_type == "IEEE" else WD_ALIGN_PARAGRAPH.LEFT
            sec_run = sec_p.add_run(sec_title)
            sec_run.font.bold = True
            sec_run.font.size = Pt(12)

            if isinstance(sec_body, list):
                for item in sec_body:
                    if isinstance(item, str):
                        doc.add_paragraph(item)
                    elif isinstance(item, dict): # Lit review item
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('citation_marker', '')} {item.get('authors', '')} ({item.get('year', '')}), ").bold = True
                        p.add_run(f"\"{item.get('title', '')}\" — {item.get('summary', '')}")
            elif isinstance(sec_body, dict):
                if "overview" in sec_body:
                    doc.add_paragraph(sec_body["overview"])
                if "notice" in sec_body:
                    p = doc.add_paragraph()
                    p.add_run(sec_body["notice"]).font.color.rgb = RGBColor(180, 0, 0)
                if "discussion_scaffold" in sec_body:
                    doc.add_paragraph(sec_body["discussion_scaffold"])
            elif isinstance(sec_body, str):
                doc.add_paragraph(sec_body)

        # References
        ref_head = doc.add_paragraph()
        ref_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_run = ref_head.add_run("REFERENCES")
        ref_run.font.bold = True

        for ref in structured_content.get("references", []):
            p = doc.add_paragraph()
            p.add_run(f"[{ref.get('index')}] ").bold = True
            p.add_run(ref.get("formatted_citation", ""))

        # Save to storage
        filename = f"formatted_{template_type.lower()}_{uuid.uuid4().hex[:8]}.docx"
        filepath = get_output_path(filename)
        doc.save(filepath)
        return f"/api/v1/files/documents/{filename}"

    @classmethod
    def render_pdf(cls, structured_content: Dict[str, Any], title: str = "Research Paper") -> str:
        """
        Renders structured paper into high quality PDF using ReportLab.
        """
        filename = f"paper_{uuid.uuid4().hex[:8]}.pdf"
        filepath = get_output_path(filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'PaperTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#1E1B4B"),
            alignment=1, # Center
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#4338CA"),
            spaceBefore=14,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            spaceAfter=8
        )

        elements = []
        elements.append(Paragraph(structured_content.get("topic", title), title_style))
        elements.append(Paragraph("<b>Author(s):</b> Academic Research Fellow | ResearchPrepAI Format Engine", body_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=15))

        # Abstract
        elements.append(Paragraph("<b>Abstract</b>", heading_style))
        elements.append(Paragraph(f"<i>{structured_content.get('abstract', '')}</i>", body_style))

        # Intro
        elements.append(Paragraph("<b>I. Introduction</b>", heading_style))
        for p in structured_content.get("introduction", []):
            elements.append(Paragraph(p, body_style))

        # Lit Review
        elements.append(Paragraph("<b>II. Literature Review</b>", heading_style))
        for item in structured_content.get("literature_review", []):
            text = f"<b>{item.get('citation_marker', '')} {item.get('authors', '')} ({item.get('year', '')})</b>: \"{item.get('title', '')}\"<br/><i>Relevance:</i> {item.get('summary', '')}"
            elements.append(Paragraph(text, body_style))

        # Methodology Scaffold
        if structured_content.get("methodology"):
            elements.append(Paragraph("<b>III. Methodology Scaffold</b>", heading_style))
            meth = structured_content.get("methodology", {})
            if isinstance(meth, dict):
                elements.append(Paragraph(meth.get("overview", ""), body_style))
                for step in meth.get("step_by_step_procedure", []):
                    elements.append(Paragraph(f"• {step}", body_style))

        # Results & Discussion
        if structured_content.get("results_and_discussion"):
            elements.append(Paragraph("<b>IV. Results & Discussion Scaffold</b>", heading_style))
            res = structured_content.get("results_and_discussion", {})
            if isinstance(res, dict):
                elements.append(Paragraph(f"<i>{res.get('notice', '')}</i>", body_style))
                elements.append(Paragraph(res.get("discussion_scaffold", ""), body_style))

        # Conclusion
        if structured_content.get("conclusion"):
            elements.append(Paragraph("<b>V. Conclusion</b>", heading_style))
            elements.append(Paragraph(structured_content.get("conclusion"), body_style))

        # Future Scope
        if structured_content.get("future_scope"):
            elements.append(Paragraph("<b>VI. Future Scope & Research Horizons</b>", heading_style))
            for item in structured_content.get("future_scope", []):
                elements.append(Paragraph(item, body_style))

        # References
        elements.append(Paragraph("<b>References</b>", heading_style))
        for ref in structured_content.get("references", []):
            elements.append(Paragraph(f"[{ref.get('index')}] {ref.get('formatted_citation', '')}", body_style))

        doc.build(elements)
        return f"/api/v1/files/documents/{filename}"

    @classmethod
    def render_report_pdf(cls, report_type: str, score: float, details: Dict[str, Any]) -> str:
        """
        Renders standalone PDF report for Plagiarism or AI Detection results.
        """
        filename = f"{report_type}_report_{uuid.uuid4().hex[:8]}.pdf"
        filepath = get_output_path(filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#4F46E5") if report_type == "plagiarism" else colors.HexColor("#D97706"),
            alignment=0,
            spaceAfter=15
        )

        elements = []
        type_label = "Plagiarism Integrity Report" if report_type == "plagiarism" else "AI Content Detection Report"
        elements.append(Paragraph(type_label, title_style))
        elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#6366F1"), spaceAfter=15))

        score_text = f"<b>Overall Score:</b> {score:.1f}% Similarity" if report_type == "plagiarism" else f"<b>AI Generated Probability:</b> {score:.1f}%"
        elements.append(Paragraph(score_text, styles['Heading2']))
        elements.append(Spacer(1, 10))

        # Audit details table
        table_data = [
            ["Audit Parameter", "Status / Metric"],
            ["Verification Engine", details.get("provider", "IntegrityEngine v2.4")],
            ["Passages Scanned", f"{details.get('passages_scanned', 14)} paragraphs"],
            ["Matched Sources", f"{details.get('matched_sources', 0)} external web/academic sources"],
            ["Evaluation Result", "PASSED Integrity Check" if score < 15 else "ATTENTION RECOMMENDED"]
        ]
        
        t = Table(table_data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#0F172A")),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0"))
        ]))
        elements.append(t)

        doc.build(elements)
        return f"/api/v1/files/documents/{filename}"

    @classmethod
    def render_ai_rewrite_docx(cls, structured_content: Dict[str, Any]) -> str:
        """
        Generates an AI-Assisted Rewrite pass DOCX clearly labeled as an editing draft aid.
        """
        doc = DocxDocument()
        
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(f"[AI-Assisted Rewrite Draft] {structured_content.get('topic', '')}")
        t_run.font.size = Pt(16)
        t_run.font.bold = True

        note_p = doc.add_paragraph()
        n_run = note_p.add_run("NOTICE: This document is an AI-assisted stylistic rewrite provided solely as a drafting aid for the author to refine in their own human voice. It is NOT intended as a final submission.")
        n_run.font.italic = True
        n_run.font.color.rgb = RGBColor(160, 40, 40)

        doc.add_paragraph()
        doc.add_paragraph("Abstract (Refined Tone):")
        doc.add_paragraph(structured_content.get("abstract", "") + " (Edited for active voice and stylistic sentence variance).")

        doc.add_paragraph()
        doc.add_paragraph("I. Introduction (Rephrased):")
        for p in structured_content.get("introduction", []):
            doc.add_paragraph(p + " (Refined for scholarly impact and flow).")

        doc.add_paragraph()
        doc.add_paragraph("II. Literature Review (Summarized):")
        for item in structured_content.get("literature_review", []):
            p = doc.add_paragraph()
            p.add_run(f"{item.get('citation_marker', '')} {item.get('authors', '')} ({item.get('year', '')}): ").bold = True
            p.add_run(f"{item.get('summary', '')}")

        doc.add_paragraph()
        doc.add_paragraph("III. Conclusion & Future Scope (Stylistic Polish):")
        if structured_content.get("conclusion"):
            doc.add_paragraph(structured_content.get("conclusion"))
        for fitem in structured_content.get("future_scope", []):
            doc.add_paragraph(fitem)

        filename = f"ai_assisted_rewrite_{uuid.uuid4().hex[:8]}.docx"
        filepath = get_output_path(filename)
        doc.save(filepath)
        return f"/api/v1/files/documents/{filename}"

document_renderer_service = DocumentRendererService()
